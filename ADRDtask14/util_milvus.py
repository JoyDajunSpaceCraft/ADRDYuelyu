import json
from collections import Counter, defaultdict
import praw
import pandas as pd
import docx


def genertate_result():
    with open("doc1_5.json", "r") as f:
        file = json.load(f)

    df = pd.read_csv("/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/embeddings.csv")
    posts = df.post.values.tolist()
    title = df.title.values.tolist()

    list1 = file["list_1"]
    list2 = file["list_2"]
    list3 = file["list_3"]
    list4 = file["list_4"]
    list5 = file["list_5"]
    for i in list5[1:6]:
        print("title :", title[i])
        print("post :", posts[i])

    return list1, list2, list3, list4, list5
    # count12 =0
    # print("overlap Doc1 & Doc 2")
    # for i in list1:
    #     if i in list2:
    #         # print("title :", title[i])
    #         # print("post :",posts[i])
    #         count12 +=1
    # print(count12/(len(list1) + len(list2) - count12))
# orders 1,2,3,4; 2,1,3,4; 3,1,2,4;

def CombMNZ():
    file_path = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/doc_idx_distance_all.json"
    with open(file_path, "r") as f:
        all_file_res = json.load(f)
    for k, v in all_file_res.items():
        if k in ["s1", "s2", "s3", "s4"]:
            min_ = 1
            max_ = 0
            for i in v:
                distance = list(i.values())[0]
                # if distance < 0.09 or distance > 1:
                #     continue
                if distance < min_:
                    min_ = distance
                if distance > max_:
                    max_ = distance

            for i in v:
                distance = list(i.values())[0]
                i[list(i.keys())[0]] = (distance - min_) / (max_ - min_)
    with open("milve_res/combMNZ.json", "w") as f:
        json.dump(all_file_res, f)


def emsamble_merge():
    # with open("milve_res/combMNZ.json", "r") as f:
    # with open("milve_res/top_50_combMNZ.json", "r") as f:
    with open("milve_res/doc_idx_distance_50.json", "r") as f:
        res = json.load(f)
    # new_res = {}
    # for docs, scores in res.items():
    #     new_sort = {}
    #     for i in scores:
    #         new_sort[list(i.keys())[0]] = list(i.values())[0]
    #     sort_scores = sorted(new_sort.items(), key=lambda scores: scores[1], reverse=True)  # from high to low
    #     count = 0
    #     new_res[docs] = []
    #     for (key, value) in sort_scores:
    #         if count < 50:
    #             new_res[docs].append({key: value})
    #             count += 1
    #         else:
    #             break
    #
    # res = new_res

    new_res = {}
    all_res = []
    # find first 20
    for k, v in res.items():
        if k in ["s1", "s2", "s3", "s4"]:
            new_res[k] = v
            for i in v:
                all_res.append(list(i.keys())[0])
    # counter = defaultdict(int)
    all_res = Counter(all_res)
    # print(all_res)
    final_rank = {}
    dict1 = {}
    dict2 = {}
    dict3 = {}
    dict4 = {}
    # change new_res to format {docid: distance}
    for k, v in new_res.items():
        if k == "s1":
            for i in v:
                dict1[list(i.keys())[0]] = list(i.values())[0]
        if k == "s2":
            for i in v:
                dict2[list(i.keys())[0]] = list(i.values())[0]
        if k == "s3":
            for i in v:
                dict3[list(i.keys())[0]] = list(i.values())[0]
        if k == "s4":
            for i in v:
                dict4[list(i.keys())[0]] = list(i.values())[0]
    # print("dict1", dict1)
    # print("dict2", dict2)
    # print("dict3", dict3)
    # print("dict4", dict4)
    all_dict = set(dict1) | set(dict2) | set(dict3) | set(dict4)
    result_3 = {}
    for m in all_dict:
        result_3[m] = dict1.get(m, 0) + dict2.get(m, 0) + dict3.get(m, 0) + dict4.get(m, 0)
    # result_3 = { # sum up all of the distance if they are in the same id
    #     key: dict1.get(key, 0) + dict2.get(key, 0) + dict3.get(key, 0) + dict4.get(key, 0) for key in
    #     set(dict1) | set(dict2) | set(dict3) | set(dict4)
    # }

    # print("result_3 is", result_3)
    # print("all_res", all_res)
    for doc_no, distance in result_3.items():
        print("begin")
        print(dict1.get(doc_no, 0))
        print(dict2.get(doc_no, 0))
        print(dict3.get(doc_no, 0))
        print(dict4.get(doc_no, 0))
        original_idx = ["409", "759", "50", "796","725"]  # orignal doc1-4 idx
        if doc_no not in original_idx:
            new_distance = distance * all_res[doc_no] # distance (CombSUM) * nonzero doc number
            final_rank[doc_no] = new_distance
            print(all_res[doc_no])
            print(new_distance)
        print("end")
    # with open("/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/final_rank.json", "w") as f:
    with open("/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/top_50_combMNZ_final_rank.json", "w") as f:
        json.dump(final_rank, f)



def match_idx_original_text():
    doc = docx.Document()
    with open("/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/final_rank.json", "r") as f:
        final_rank = json.load(f)
    final_rank = {k: v for k, v in sorted(final_rank.items(), key=lambda item: item[1],reverse=True)}
    # print(final_rank)
    files2docid = match_original()
    all_df = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/new_all_adrd.csv"
    df = pd.read_csv(all_df)
    client_id = "pV-DoxK3yzaD8eJTuMBRWQ"
    client_secret = "dQvLvNze_-hu6DRhIU_Nsu6Od12E3A"
    username = "yuj49"
    password = "Fiona@1999"
    user_agent = "yuj49"
    reddit = praw.Reddit(client_id=client_id,
                         client_secret=client_secret,
                         username=username,
                         password=password,
                         user_agent=user_agent)
    count = 1
    for k, v in final_rank.items():
        if count > 50:
            break
        id = df.loc[int(k)].id
        res = files2docid[k]
        items = []
        # avgs = 0
        for i in res:
            # avgs += i[1]
            i[1] = str(i[1])
            items.append(": ".join(i))


        if reddit.submission(id=id).title not in ["[deleted]", "[removed]", "[deleted by user]"] and reddit.submission(
                id=id).selftext not in ["[deleted]", "[removed]", "[deleted by user]"]:
            doc.add_heading(
                "search result id is {} from related doc {} and average distance is {}".format(count, ", ".join(items),
                                                                                               avgs), 1)
            doc.add_heading("Title: " + reddit.submission(id=id).title, 2)
            doc.add_heading("Unique ID: " + id, 3)
            doc.add_paragraph("Post: " + reddit.submission(id=id).selftext)

            # Adding a page break
            doc.add_page_break()
            count += 1
        # except Exception as e:
        #     print("cannot find ",id)
    doc.save('50_search_result.docx')
    # print(reddit.info(fullnames=[id]))
    # print(df.loc[int(k)].id)


def match_original():
    with open("milve_res/combMNZ.json", "r") as f:
        res = json.load(f)
    with open("/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/final_rank.json", "r") as f:
        final_rank = json.load(f)
    ids = list(final_rank.keys())
    res_ = {"s1": [], "s2": [], "s3": [], "s4": []}
    for k, v in res.items():
        if k in ["s1", "s2", "s3", "s4"]:
            for i in v:
                res_[k].append([list(i.keys())[0], list(i.values())[0]])
    doc2index = {}
    for k, v in res_.items():
        # find the k in the
        if k == "s1":
            k = "1"
        if k == "s2":
            k = "2"
        if k == "s3":
            k = "3"
        if k == "s4":
            k = "4"
        for i in v:

            if i[0] in ids and i[0] not in doc2index.keys():
                doc2index[i[0]] = [[k, i[1]]]
            elif i[0] in ids and i[0] in doc2index.keys():
                doc2index[i[0]].append([k, i[1]])
                doc2index[i[0]] = sorted(doc2index[i[0]], key=lambda docids: docids[1], reverse=True)
    # print("original dox2idx",doc2index)
    return doc2index


def print_result():
    # path = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/final_rank.json"
    path = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/doc_idx_distance_all.json"
    original = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/embeddings.csv"

    df = pd.read_csv(original)
    post = df.post.values.tolist()
    title = df.title.values.tolist()

    with open(path, "r") as f:
        j_ = json.load(f)
    file1234 = j_["order1234"]
    file4321 = j_["order4321"]

    # print("the result of the emsemble result")
    print("here is the result of the 1234")
    for i in file1234[1:6]:
        print(list(i.keys())[0])
        idx = int(list(i.keys())[0])
        print("title: ", title[idx])
        print("post: ", post[idx])

    print("******************here is the result of the 4321")
    for i in file4321[:5]:
        print(list(i.keys())[0])
        idx = int(list(i.keys())[0])
        print("title: ", title[idx])
        print("post: ", post[idx])


def get_top50_threshold():
    with open("milve_res/combMNZ.json", "r") as f:
        res = json.load(f)
    threshold = {}
    for docs, scores in res.items():
        new_sort = {}
        for i in scores:
            new_sort[list(i.keys())[0]] = list(i.values())[0]
        sort_scores = sorted(new_sort.items(), key=lambda scores: scores[1], reverse=True)  # from high to low
        count = 0
        thresholds = 0
        for (key, values) in sort_scores:
            if count==50:
                thresholds = values
                break
            count+=1
        threshold[docs] = thresholds
    print(threshold) # {'s1': 0.36127701088542735, 's2': 0.3548173091785063, 's3': 0.2511407754987297, 's4': 0.30168559377768445}

def generate_top50_combinz(): # generate top 50
    file_path = "/Users/yuelyu/PycharmProjects/ADRD/ADRDtask14/milve_res/doc_idx_distance_all.json"

    with open(file_path, "r") as f:
        all_file_res = json.load(f)

    new_top_50 = {}
    new_all_file_res = {}

    for k, v in all_file_res.items():
        all_new_sort = {}
        for i in v:
            all_new_sort[list(i.keys())[0]] = list(i.values())[0]
        all_sort_scores = sorted(all_new_sort.items(), key=lambda scores: scores[1], reverse=True)  # from high to low
        new_all_file_res[k] = []
        count = 0
        for (key,value) in all_sort_scores:
            if count <50:
                new_all_file_res[k].append({key:value})
                count+=1
            else:
                break

    for k, v in new_all_file_res.items():
        if k in ["s1", "s2", "s3", "s4"]:
            min_ = 1
            max_ = 0
            for i in v:
                distance = list(i.values())[0]
                # if distance < 0.09 or distance > 1:
                #     continue
                if distance < min_:
                    min_ = distance
                if distance > max_:
                    max_ = distance
            for i in v:
                distance = list(i.values())[0]
                i[list(i.keys())[0]] = (distance - min_) / (max_ - min_)
        new_top_50[k] = []
        new_sort = {}
        for i in v:
            new_sort[list(i.keys())[0]] = list(i.values())[0]
        sort_scores = sorted(new_sort.items(), key=lambda scores: scores[1], reverse=True)  # from high to low
        count = 0

        for (key, value) in sort_scores:
            if count<50:
                count+=1
                new_top_50[k].append({key:value})
            else:
                break

    with open("milve_res/top_50_combMNZ.json", "w") as f:
        json.dump(new_top_50, f)

if __name__ == "__main__":
    # CombMNZ()
    emsamble_merge()
    # print_result()
    # match_original()
    # match_idx_original_text()
    # get_top50_threshold()
    # generate_top50_combinz()
